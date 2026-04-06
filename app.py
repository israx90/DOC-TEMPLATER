# -*- coding: utf-8 -*-
from flask import Flask, render_template, request, jsonify, send_file, Response, stream_with_context
import os
import sys
import json
import io
import zipfile
import shutil
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Frozen (PyInstaller) path detection ---
FROZEN = getattr(sys, 'frozen', False)
if FROZEN:
    # PyInstaller bundles static/templates into sys._MEIPASS
    BUNDLE_DIR = sys._MEIPASS
    # Writable dirs go to ~/Documents/EDTECH DOC TEMPLATER/
    USER_DATA = os.path.join(os.path.expanduser('~'), 'Documents', 'EDTECH DOC TEMPLATER')
    os.makedirs(USER_DATA, exist_ok=True)
else:
    BUNDLE_DIR = os.path.dirname(os.path.abspath(__file__))
    USER_DATA = os.path.dirname(os.path.abspath(__file__))

# Optional OCR imports (fail gracefully if not installed)
try:
    from img2table.ocr import TesseractOCR
    from img2table.document import Image as Img2TableImage
    import tempfile
    # Configure Tesseract path for frozen builds
    if FROZEN:
        import pytesseract
        _tess_dir = os.path.join(BUNDLE_DIR, 'tesseract')
        _tess_bin = os.path.join(_tess_dir, 'tesseract')
        if os.path.exists(_tess_bin):
            pytesseract.pytesseract.tesseract_cmd = _tess_bin
            os.environ['TESSDATA_PREFIX'] = os.path.join(_tess_dir, 'tessdata')
            # Set library path so dyld can find bundled dylibs
            os.environ['DYLD_LIBRARY_PATH'] = _tess_dir + ':' + os.environ.get('DYLD_LIBRARY_PATH', '')
            os.environ['DYLD_FALLBACK_LIBRARY_PATH'] = _tess_dir
    HAS_OCR = True
except Exception as _ocr_err:
    HAS_OCR = False
    print('OCR WARNING: img2table not available - {}: {}'.format(type(_ocr_err).__name__, _ocr_err))

# Monkey-patch python-docx template path for frozen builds
if FROZEN:
    try:
        _docx_tpl_dir = os.path.join(BUNDLE_DIR, 'docx', 'templates')
        if os.path.isdir(_docx_tpl_dir):
            import docx.parts.hdrftr as _hdrftr

            @classmethod
            def _patched_header_xml(cls):
                path = os.path.join(_docx_tpl_dir, 'default-header.xml')
                with open(path, 'rb') as f:
                    return f.read()

            @classmethod
            def _patched_footer_xml(cls):
                path = os.path.join(_docx_tpl_dir, 'default-footer.xml')
                with open(path, 'rb') as f:
                    return f.read()

            _hdrftr.HeaderPart._default_header_xml = _patched_header_xml
            _hdrftr.FooterPart._default_footer_xml = _patched_footer_xml
            print('[PATCH] python-docx templates patched to: {}'.format(_docx_tpl_dir))
        else:
            print('[PATCH WARNING] docx/templates dir not found at: {}'.format(_docx_tpl_dir))
    except Exception as _patch_err:
        print('[PATCH ERROR] Failed to patch python-docx: {}'.format(_patch_err))

app = Flask(__name__,
            static_folder=os.path.join(BUNDLE_DIR, 'static'),
            template_folder=os.path.join(BUNDLE_DIR, 'templates'))

# Configuration
UPLOAD_FOLDER = os.path.join(USER_DATA, 'uploads')
OUTPUT_FOLDER = os.path.join(USER_DATA, 'outputs')

# Clean previous session files on startup
for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
    if os.path.exists(folder):
        shutil.rmtree(folder)
    os.makedirs(folder, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

# --- Routes ---

@app.route('/')
def index():
    return render_template('tool_docs.html')

@app.route('/docs')
def docs():
    return render_template('tool_docs.html')

@app.route('/debug_info')
def debug_info():
    """Return diagnostic info for debugging on other Macs."""
    import platform, subprocess, sys
    info = {
        'frozen': FROZEN,
        'bundle_dir': BUNDLE_DIR,
        'user_data': USER_DATA,
        'upload_folder': UPLOAD_FOLDER,
        'output_folder': OUTPUT_FOLDER,
        'python_version': sys.version,
        'platform': platform.platform(),
        'arch': platform.machine(),
        'has_ocr': HAS_OCR,
        'tesseract': {'available': False, 'version': None, 'error': None},
        'static_folder': app.static_folder,
        'template_folder': app.template_folder,
    }
    info['dyld_library_path'] = os.environ.get('DYLD_LIBRARY_PATH', 'NOT SET')
    info['dyld_fallback'] = os.environ.get('DYLD_FALLBACK_LIBRARY_PATH', 'NOT SET')
    # Check Tesseract
    candidates = []
    tess_env = os.environ.copy()
    if FROZEN:
        bundled = os.path.join(BUNDLE_DIR, 'tesseract', 'tesseract')
        tess_dir = os.path.join(BUNDLE_DIR, 'tesseract')
        info['tesseract']['bundled_path'] = bundled
        info['tesseract']['bundled_exists'] = os.path.exists(bundled)
        # List bundled dylibs
        if os.path.isdir(tess_dir):
            info['tesseract']['bundled_files'] = [f for f in os.listdir(tess_dir) if f.endswith('.dylib')]
        tess_env['DYLD_LIBRARY_PATH'] = tess_dir
        tess_env['DYLD_FALLBACK_LIBRARY_PATH'] = tess_dir
        candidates.append(bundled)
    candidates.append('tesseract')
    for tess_cmd in candidates:
        try:
            result = subprocess.run([tess_cmd, '--version'], capture_output=True, text=True, timeout=5, env=tess_env)
            info['tesseract']['cmd'] = tess_cmd
            info['tesseract']['returncode'] = result.returncode
            info['tesseract']['stdout'] = result.stdout[:200] if result.stdout else ''
            info['tesseract']['stderr'] = result.stderr[:200] if result.stderr else ''
            if result.returncode == 0:
                info['tesseract']['available'] = True
                version_out = result.stdout or result.stderr or ''
                info['tesseract']['version'] = version_out.strip().split('\n')[0]
                break
        except Exception as e:
            info['tesseract']['error'] = '{}: {}'.format(type(e).__name__, str(e))
    return jsonify(info)

@app.route('/check_tesseract')
def check_tesseract():
    """Check if Tesseract OCR is available."""
    import subprocess
    # Try bundled binary first (frozen), then system PATH
    candidates = []
    if FROZEN:
        candidates.append(os.path.join(BUNDLE_DIR, 'tesseract', 'tesseract'))
    candidates.append('tesseract')  # system PATH fallback

    tess_env = os.environ.copy()
    if FROZEN:
        tess_dir = os.path.join(BUNDLE_DIR, 'tesseract')
        tess_env['DYLD_LIBRARY_PATH'] = tess_dir
        tess_env['DYLD_FALLBACK_LIBRARY_PATH'] = tess_dir

    for tess_cmd in candidates:
        try:
            result = subprocess.run([tess_cmd, '--version'], capture_output=True, text=True, timeout=5, env=tess_env)
            if result.returncode != 0:
                continue  # binary crashed (e.g. dyld symbol error), try next
            version_output = result.stdout or result.stderr or ''
            version = version_output.strip().split('\n')[0] if version_output.strip() else 'unknown'
            return jsonify({'available': True, 'version': version, 'has_ocr_deps': HAS_OCR})
        except Exception:
            continue  # FileNotFoundError, TimeoutExpired, etc.

    return jsonify({'available': False, 'version': None, 'has_ocr_deps': HAS_OCR})

@app.route('/upload_header', methods=['POST'])
def upload_header():
    if 'headerUtils' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['headerUtils']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        upload_folder = app.config['UPLOAD_FOLDER']
        for existing in os.listdir(upload_folder):
            if existing.startswith('custom_header.'):
                os.remove(os.path.join(upload_folder, existing))
        ext = os.path.splitext(file.filename)[1].lower() or '.png'
        filename = 'custom_header{}'.format(ext)
        filepath = os.path.join(upload_folder, filename)
        file.save(filepath)
        return jsonify({'path': '/uploads/{}'.format(filename), 'success': True, 'filename': filename})

@app.route('/upload_footer', methods=['POST'])
def upload_footer():
    if 'footerUtils' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['footerUtils']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        upload_folder = app.config['UPLOAD_FOLDER']
        for existing in os.listdir(upload_folder):
            if existing.startswith('custom_footer.'):
                os.remove(os.path.join(upload_folder, existing))
        ext = os.path.splitext(file.filename)[1].lower() or '.png'
        filename = 'custom_footer{}'.format(ext)
        filepath = os.path.join(upload_folder, filename)
        file.save(filepath)
        return jsonify({'path': '/uploads/{}'.format(filename), 'success': True, 'filename': filename})

@app.route('/upload_cover', methods=['POST'])
def upload_cover():
    if 'coverUtils' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['coverUtils']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        upload_folder = app.config['UPLOAD_FOLDER']
        for existing in os.listdir(upload_folder):
            if existing.startswith('custom_cover.'):
                os.remove(os.path.join(upload_folder, existing))
        ext = os.path.splitext(file.filename)[1].lower() or '.png'
        filename = 'custom_cover{}'.format(ext)
        filepath = os.path.join(upload_folder, filename)
        file.save(filepath)
        return jsonify({'path': '/uploads/{}'.format(filename), 'success': True, 'filename': filename})

@app.route('/upload_backpage', methods=['POST'])
def upload_backpage():
    if 'backpageUtils' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['backpageUtils']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file:
        upload_folder = app.config['UPLOAD_FOLDER']
        for existing in os.listdir(upload_folder):
            if existing.startswith('custom_backpage.'):
                os.remove(os.path.join(upload_folder, existing))
        ext = os.path.splitext(file.filename)[1].lower() or '.png'
        filename = 'custom_backpage{}'.format(ext)
        filepath = os.path.join(upload_folder, filename)
        file.save(filepath)
        return jsonify({'path': '/uploads/{}'.format(filename), 'success': True, 'filename': filename})

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename))


@app.route('/download/<path:filename>')
def download_file(filename):
    try:
        return send_file(os.path.join(app.config['OUTPUT_FOLDER'], filename), as_attachment=True)
    except Exception as e:
        return str(e), 404


# ─────────────────────────────────────────────────────────────────────────────
# .EDD TEMPLATE FORMAT  (EDTech Document)
# A ZIP archive containing:
#   manifest.json  — version metadata
#   config.json    — all style/format settings
#   images/        — cover, header, footer, backpage (if present)
# ─────────────────────────────────────────────────────────────────────────────

@app.route('/save_template', methods=['POST'])
def save_template():
    """Package current config + images into a .edd file (ZIP-based)."""
    import json as _json

    # Config comes as JSON body
    try:
        cfg = request.get_json(force=True) or {}
    except Exception:
        cfg = {}

    upload_folder = app.config['UPLOAD_FOLDER']
    image_slots = {
        'cover':    'custom_cover',
        'header':   'custom_header',
        'footer':   'custom_footer',
        'backpage': 'custom_backpage',
    }
    image_exts = ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.emf', '.wmf']

    # Build manifest
    import datetime
    manifest = {
        'format': 'edd',
        'version': '1.0',
        'app': 'EDTech Doc Templater',
        'created': datetime.datetime.utcnow().isoformat() + 'Z',
    }

    zip_buffer = io.BytesIO()
    images_included = {}

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('manifest.json', _json.dumps(manifest, indent=2))
        zf.writestr('config.json', _json.dumps(cfg, indent=2, ensure_ascii=False))

        for slot, base in image_slots.items():
            for ext in image_exts:
                candidate = os.path.join(upload_folder, base + ext)
                if os.path.exists(candidate):
                    arcname = 'images/{}{}'.format(slot, ext)
                    zf.write(candidate, arcname)
                    images_included[slot] = arcname
                    break  # only first match per slot

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        mimetype='application/octet-stream',
        as_attachment=True,
        download_name='plantilla.edd'
    )


@app.route('/load_template', methods=['POST'])
def load_template():
    """Load a .edd file: restore images to uploads folder, return config JSON."""
    import json as _json

    edd_file = request.files.get('edd')
    if not edd_file:
        return jsonify({'error': 'No se recibió ningún archivo .edd'}), 400

    try:
        zf = zipfile.ZipFile(io.BytesIO(edd_file.read()))
    except zipfile.BadZipFile:
        return jsonify({'error': 'El archivo .edd está corrupto o no es válido'}), 400

    names = zf.namelist()

    # Validate structure
    if 'config.json' not in names:
        return jsonify({'error': 'Archivo .edd inválido: falta config.json'}), 400

    # Read config
    try:
        cfg = _json.loads(zf.read('config.json').decode('utf-8'))
    except Exception as e:
        return jsonify({'error': 'Error al leer config.json: {}'.format(e)}), 400

    upload_folder = app.config['UPLOAD_FOLDER']
    slot_map = {
        'cover':    'custom_cover',
        'header':   'custom_header',
        'footer':   'custom_footer',
        'backpage': 'custom_backpage',
    }

    restored_images = {}  # slot -> URL path the client can preview

    for name in names:
        if not name.startswith('images/'):
            continue
        basename = os.path.basename(name)   # e.g. "cover.png"
        slot, ext = os.path.splitext(basename)  # "cover", ".png"
        if slot not in slot_map:
            continue

        # Remove any old image for this slot
        for old_ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.emf', '.wmf']:
            old = os.path.join(upload_folder, slot_map[slot] + old_ext)
            if os.path.exists(old):
                os.remove(old)

        dest_name = slot_map[slot] + ext
        dest_path = os.path.join(upload_folder, dest_name)
        with open(dest_path, 'wb') as f:
            f.write(zf.read(name))

        # Return a URL the browser can preview (reuse existing /uploads/<path> pattern)
        restored_images[slot] = '/uploads/{}'.format(dest_name)

    return jsonify({'config': cfg, 'images': restored_images})


@app.route('/process', methods=['POST'])
def process_docs():
    files = request.files.getlist('docs')
    output_folder_path = request.form.get('output_folder', app.config['OUTPUT_FOLDER'])
    paper_size = request.form.get('paper_size', 'letter')
    prefix = request.form.get('prefix', 'FORMAT_')
    suffix = request.form.get('suffix', '')
    export_format = request.form.get('export_format', 'docx')  # docx, pdf, both

    if output_folder_path:
        try:
            os.makedirs(output_folder_path, exist_ok=True)
            target_folder = output_folder_path
        except Exception:
            target_folder = app.config['OUTPUT_FOLDER']
    else:
        target_folder = app.config['OUTPUT_FOLDER']

    style_config_str = request.form.get('style_config', '{}')
    try:
        style_config = json.loads(style_config_str)
    except Exception:
        style_config = {}

    # Save all files to disk BEFORE the generator runs (request context closes after response starts)
    saved_files = []
    for file in files:
        original_name = file.filename
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        saved_files.append((original_name, filename, input_path))

    def generate():
        import time as _time
        t_start = _time.time()
        results = []
        total_pages = 0

        # Check PDF conversion availability if needed
        pdf_converter = None
        if export_format in ('pdf', 'both'):
            try:
                from docx2pdf import convert as docx2pdf_convert
                pdf_converter = docx2pdf_convert
            except ImportError:
                yield json.dumps({'msg': '⚠️ docx2pdf no instalado — exportando solo DOCX. Ejecuta: pip3 install docx2pdf', 'type': 'error'}) + '\n'

        for original_name, filename, input_path in saved_files:
            yield json.dumps({'msg': 'Transformando: {}...'.format(original_name), 'type': 'info'}) + '\n'

            try:
                # Step 1: Load document
                yield json.dumps({'msg': '[DEBUG] Cargando: {}'.format(input_path), 'type': 'info'}) + '\n'
                doc = Document(input_path)
                page_count = len(doc.sections)
                yield json.dumps({'msg': '[DEBUG] Cargado OK: {} secciones'.format(page_count), 'type': 'info'}) + '\n'

                # Step 2: Apply styles
                yield json.dumps({'msg': '[DEBUG] Aplicando estilos...', 'type': 'info'}) + '\n'
                apply_styles(doc, style_config, paper_size)
                yield json.dumps({'msg': '[DEBUG] Estilos aplicados OK', 'type': 'info'}) + '\n'

                # Step 3: Build output path
                base_name = os.path.splitext(filename)[0]
                output_base = '{}{}{}'.format(prefix, base_name, suffix)
                output_docx_name = '{}.docx'.format(output_base)
                output_docx_path = os.path.join(target_folder, output_docx_name)
                yield json.dumps({'msg': '[DEBUG] Guardando en: {}'.format(output_docx_path), 'type': 'info'}) + '\n'

                # Step 4: Save
                doc.save(output_docx_path)
                yield json.dumps({'msg': '[DEBUG] Guardado OK: {}'.format(output_docx_name), 'type': 'info'}) + '\n'

                pdf_name = None

                # PDF conversion (if requested and available)
                if pdf_converter and export_format in ('pdf', 'both'):
                    try:
                        pdf_output_name = '{}.pdf'.format(output_base)
                        pdf_output_path = os.path.join(target_folder, pdf_output_name)
                        pdf_converter(output_docx_path, pdf_output_path)
                        pdf_name = pdf_output_name
                        yield json.dumps({'msg': 'PDF generado: {}'.format(pdf_output_name), 'type': 'info'}) + '\n'

                        # If only PDF requested, remove the intermediate DOCX
                        if export_format == 'pdf':
                            os.remove(output_docx_path)
                            output_docx_name = None
                    except Exception as pdf_err:
                        yield json.dumps({'msg': 'Error PDF {}: {}'.format(original_name, str(pdf_err)), 'type': 'error'}) + '\n'

                total_pages += page_count
                results.append({
                    'name': original_name,
                    'status': 'success',
                    'docx': output_docx_name,
                    'pdf': pdf_name,
                    'pages': page_count
                })
                yield json.dumps({'msg': 'Completado: {}'.format(original_name), 'type': 'success'}) + '\n'
            except Exception as e:
                import traceback
                tb = traceback.format_exc()
                print('[EXPORT ERROR] {}: {}'.format(original_name, tb))
                yield json.dumps({'msg': '[ERROR] {} → {}: {}'.format(original_name, type(e).__name__, str(e)), 'type': 'error'}) + '\n'
                yield json.dumps({'msg': '[TRACEBACK] {}'.format(tb.replace('\n', ' | ')), 'type': 'error'}) + '\n'
                results.append({'name': original_name, 'status': 'error', 'error': str(e)})

        elapsed = _time.time() - t_start
        stats = {
            'total_files': len(results),
            'success_count': sum(1 for r in results if r['status'] == 'success'),
            'error_count': sum(1 for r in results if r['status'] == 'error'),
            'total_pages': total_pages,
            'elapsed_seconds': round(elapsed, 1),
        }
        print('[EXPORT] Stats: success={}, errors={}, total={}, folder={}'.format(
            stats['success_count'], stats['error_count'], stats['total_files'], target_folder))
        complete_json = json.dumps({'type': 'complete', 'results': results, 'output_folder': target_folder, 'stats': stats})
        print('[EXPORT] Complete JSON length: {} bytes'.format(len(complete_json)))
        yield complete_json + '\n'

    return Response(stream_with_context(generate()), mimetype='application/json')


# --- HELPER FUNCTIONS ---

def hex_to_rgb(hex_code):
    hex_code = hex_code.lstrip('#')
    return tuple(int(hex_code[i:i+2], 16) for i in (0, 2, 4))

def set_cell_border(cell, **kwargs):
    """Set borders on a table cell. Pass top/bottom/left/right as dicts with val, sz, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement('w:{}'.format(edge))
            for attr_name, attr_val in kwargs[edge].items():
                element.set(qn('w:{}'.format(attr_name)), attr_val)
            tcBorders.append(element)
    tcPr.append(tcBorders)

def border_style_to_val(style):
    mapping = {'single': 'single', 'thick': 'thick', 'dashed': 'dashed', 'none': 'nil'}
    return mapping.get(style, 'single')

def embed_header_image(doc):
    """Embed the uploaded header image (PNG/JPG/EMF/WMF) into every section's header."""
    upload_folder = app.config['UPLOAD_FOLDER']

    # Find any custom_header file (any extension)
    header_path = None
    header_ext = None
    for ext in ('.png', '.jpg', '.jpeg', '.emf', '.wmf', '.gif', '.bmp'):
        candidate = os.path.join(upload_folder, 'custom_header{}'.format(ext))
        if os.path.exists(candidate):
            header_path = candidate
            header_ext = ext.lower()
            break

    if not header_path:
        return  # No header uploaded, skip

    is_vector = header_ext in ('.emf', '.wmf')

    try:
        for section in doc.sections:
            # Unlink so header shows its own content
            section.header.is_linked_to_previous = False
            section.header_distance = Cm(0)
            header = section.header

            page_width = section.page_width or Inches(8.5)

            # Clear all existing header paragraph content
            for para in header.paragraphs:
                p = para._p
                for child in list(p):
                    p.remove(child)

            # Use first paragraph or add one (will be an empty anchor container)
            if header.paragraphs:
                para = header.paragraphs[0]
            else:
                para = header.add_paragraph()

            # No indent needed — behind-text images are positioned absolutely
            pf = para.paragraph_format
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

            if not is_vector:
                _embed_raster_behind_text(para, header_path, header_ext, width_emu=page_width)
            else:
                _embed_vector_image(para, header_path, header_ext, width_emu=page_width)

    except Exception as e:
        print('Header embedding error: {}'.format(e))


def clean_footers(doc):
    """Remove all content from footers in all sections (prepare for clean insertion)."""
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        # Remove all paragraphs securely
        for para in list(footer.paragraphs):
            p = para._p
            if p.getparent() is not None:
                p.getparent().remove(p)

def embed_footer_image(doc):
    """Embed the uploaded footer image (PNG/JPG/EMF/WMF) into every section's footer (behind text)."""
    upload_folder = app.config['UPLOAD_FOLDER']
    footer_path = None
    footer_ext = None
    for ext in ('.png', '.jpg', '.jpeg', '.emf', '.wmf', '.gif', '.bmp'):
        candidate = os.path.join(upload_folder, 'custom_footer{}'.format(ext))
        if os.path.exists(candidate):
            footer_path = candidate
            footer_ext = ext.lower()
            break

    if not footer_path:
        return

    is_vector = footer_ext in ('.emf', '.wmf')

    try:
        for section in doc.sections:
            # Note: We assume clean_footers(doc) was called, so footer is empty or ready.
            section.footer.is_linked_to_previous = False
            # 1.27cm = 0.5 inch (standard). 0 would be edge of page.
            section.footer_distance = Cm(1.27)
            footer = section.footer

            page_width = section.page_width or Inches(8.5)
            page_height = section.page_height or Inches(11)

            # Add a new paragraph for the image anchor
            para = footer.add_paragraph()

            pf = para.paragraph_format
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

            if not is_vector:
                _embed_raster_footer_behind_text(para, footer_path, footer_ext,
                                                 width_emu=page_width,
                                                 page_height_emu=page_height)
            else:
                _embed_vector_footer(para, footer_path, footer_ext,
                                     width_emu=page_width, page_height_emu=page_height)

    except Exception as e:
        print('Footer embedding error: {}'.format(e))


def _embed_raster_footer_behind_text(para, image_path, ext, width_emu=None, page_height_emu=None):
    """Embed PNG/JPG as a behind-text anchor image pinned to the page bottom."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    import lxml.etree as etree
    import uuid as _uuid

    content_type_map = {
        '.png': 'image/png', '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg', '.gif': 'image/gif', '.bmp': 'image/bmp',
    }
    content_type = content_type_map.get(ext.lower(), 'image/png')

    with open(image_path, 'rb') as f:
        image_bytes = f.read()

    try:
        from PIL import Image as _PIL_Image
        import io as _io
        img = _PIL_Image.open(_io.BytesIO(image_bytes))
        img_w, img_h = img.size
        if width_emu is None:
            width_emu = int(7.5 * 914400)
        height_emu = int(width_emu * img_h / img_w)
    except Exception:
        if width_emu is None:
            width_emu = int(7.5 * 914400)
        height_emu = int(914400)

    # Position: bottom of page — posOffset = page_height - image_height
    if page_height_emu is None:
        page_height_emu = int(11 * 914400)
    pos_y = int(page_height_emu) - int(height_emu)

    footer_part = para.part
    uid = _uuid.uuid4().hex[:8]
    part_name = PackURI('/word/media/ftr_{}{}'.format(uid, ext))
    image_part = Part(part_name, content_type, image_bytes)
    rId = footer_part.relate_to(
        image_part,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    )
    img_id = (abs(hash(uid)) % 65000) + 200

    anchor_xml = (
        '<w:drawing'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        ' relativeHeight="2" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>{pos_y}</wp:posOffset></wp:positionV>'
        '<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="{img_id}" name="FtrImg{img_id}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic>'
        '<pic:nvPicPr>'
        '<pic:cNvPr id="{img_id}" name="FtrImg{img_id}"/>'
        '<pic:cNvPicPr/>'
        '</pic:nvPicPr>'
        '<pic:blipFill>'
        '<a:blip r:embed="{rId}"/>'
        '<a:stretch><a:fillRect/></a:stretch>'
        '</pic:blipFill>'
        '<pic:spPr>'
        '<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '</pic:spPr>'
        '</pic:pic>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:anchor>'
        '</w:drawing>'
    ).format(cx=int(width_emu), cy=int(height_emu), pos_y=pos_y, rId=rId, img_id=img_id)

    drawing_element = etree.fromstring(anchor_xml)
    p = para._p
    r_elem = OxmlElement('w:r')
    r_elem.append(drawing_element)
    p.append(r_elem)


def _embed_vector_footer(para, image_path, ext, width_emu=None, page_height_emu=None):
    """Embed WMF/EMF as behind-text VML in footer, anchored to page bottom."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    import lxml.etree as etree
    import uuid as _uuid

    content_type_map = {'.emf': 'image/x-emf', '.wmf': 'image/x-wmf'}
    content_type = content_type_map.get(ext, 'image/x-wmf')

    with open(image_path, 'rb') as f:
        image_bytes = f.read()

    footer_part = para.part
    part_name = PackURI('/word/media/ftr_{}{}'.format(_uuid.uuid4().hex[:8], ext))
    image_part = Part(part_name, content_type, image_bytes)
    rId = footer_part.relate_to(
        image_part,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    )

    if width_emu is None:
        width_emu = int(7.5 * 914400)
    if page_height_emu is None:
        page_height_emu = int(11 * 914400)
    w_pt = int(width_emu / 12700)
    h_pt = int(1.0 * 914400 / 12700)
    # margin-top calculated so it sits at the bottom
    mt_pt = int((page_height_emu - int(914400)) / 12700)

    pict_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:v="urn:schemas-microsoft-com:vml"'
        ' xmlns:o="urn:schemas-microsoft-com:office:office"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<v:shape style="position:absolute;margin-left:0pt;margin-top:{mt}pt;'
        'width:{w}pt;height:{h}pt;z-index:-251658240;'
        'mso-position-horizontal-relative:margin;mso-position-vertical-relative:margin"'
        ' filled="f" stroked="f">'
        '<v:imagedata r:id="{rId}" o:title="footer"/>'
        '</v:shape>'
        '</w:pict>'
    ).format(w=w_pt, h=h_pt, mt=mt_pt, rId=rId)

    pict_element = etree.fromstring(pict_xml)
    p = para._p
    r_elem = OxmlElement('w:r')
    r_elem.append(pict_element)
    p.append(r_elem)



def _embed_raster_behind_text(para, image_path, ext, width_emu=None):
    """Embed PNG/JPG as a behind-text floating anchor image (behindDoc=1) in a header paragraph."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    import lxml.etree as etree
    import uuid as _uuid

    content_type_map = {
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
    }
    content_type = content_type_map.get(ext.lower(), 'image/png')

    with open(image_path, 'rb') as f:
        image_bytes = f.read()

    # Calculate proportional height from actual image dimensions
    try:
        from PIL import Image as _PIL_Image
        import io as _io
        img = _PIL_Image.open(_io.BytesIO(image_bytes))
        img_w, img_h = img.size
        if width_emu is None:
            width_emu = int(7.5 * 914400)
        height_emu = int(width_emu * img_h / img_w)
    except Exception:
        if width_emu is None:
            width_emu = int(7.5 * 914400)
        height_emu = int(914400)  # fallback: 1 inch

    header_part = para.part
    uid = _uuid.uuid4().hex[:8]
    part_name = PackURI('/word/media/hdr_{}{}'.format(uid, ext))
    image_part = Part(part_name, content_type, image_bytes)
    rId = header_part.relate_to(
        image_part,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    )

    img_id = (abs(hash(uid)) % 65000) + 100

    # wp:anchor with behindDoc="1" — positions image absolutely at (0,0) from page edge,
    # behind all text content
    anchor_xml = (
        '<w:drawing'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        ' relativeHeight="2" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        '<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="{img_id}" name="HdrImg{img_id}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic>'
        '<pic:nvPicPr>'
        '<pic:cNvPr id="{img_id}" name="HdrImg{img_id}"/>'
        '<pic:cNvPicPr/>'
        '</pic:nvPicPr>'
        '<pic:blipFill>'
        '<a:blip r:embed="{rId}"/>'
        '<a:stretch><a:fillRect/></a:stretch>'
        '</pic:blipFill>'
        '<pic:spPr>'
        '<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '</pic:spPr>'
        '</pic:pic>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:anchor>'
        '</w:drawing>'
    ).format(cx=int(width_emu), cy=int(height_emu), rId=rId, img_id=img_id)

    drawing_element = etree.fromstring(anchor_xml)
    p = para._p
    r_elem = OxmlElement('w:r')
    r_elem.append(drawing_element)
    p.append(r_elem)


def _embed_vector_image(para, image_path, ext, width_emu=None):
    """Embed WMF/EMF as a behind-text floating VML image (position:absolute, z-index behind)."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    import lxml.etree as etree
    import uuid as _uuid

    content_type_map = {
        '.emf': 'image/x-emf',
        '.wmf': 'image/x-wmf',
    }
    content_type = content_type_map.get(ext, 'image/x-wmf')

    with open(image_path, 'rb') as f:
        image_bytes = f.read()

    header_part = para.part
    part_name = PackURI('/word/media/hdr_{}{}'.format(_uuid.uuid4().hex[:8], ext))
    image_part = Part(part_name, content_type, image_bytes)
    rId = header_part.relate_to(
        image_part,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    )

    if width_emu is None:
        width_emu = int(7.5 * 914400)
    w_pt = int(width_emu / 12700)
    h_pt = int(1.0 * 914400 / 12700)  # 1 inch tall

    # position:absolute + z-index:-251658240 = Word's "Behind Text" for VML shapes
    pict_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:v="urn:schemas-microsoft-com:vml"'
        ' xmlns:o="urn:schemas-microsoft-com:office:office"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<v:shape style="position:absolute;margin-left:0pt;margin-top:0pt;'
        'width:{w}pt;height:{h}pt;z-index:-251658240;'
        'mso-position-horizontal-relative:margin;mso-position-vertical-relative:margin"'
        ' filled="f" stroked="f">'
        '<v:imagedata r:id="{rId}" o:title="header"/>'
        '</v:shape>'
        '</w:pict>'
    ).format(w=w_pt, h=h_pt, rId=rId)

    pict_element = etree.fromstring(pict_xml)
    p = para._p
    r_elem = OxmlElement('w:r')
    r_elem.append(pict_element)
    p.append(r_elem)


def insert_toc_page(doc, config):
    """
    Insert a Table of Contents page after the cover page.
    Uses Word's TOC field code — headings must have Heading 1/2 styles assigned.
    The updateFields setting ensures Word auto-populates the TOC on open.
    """
    toc_cfg = config.get('toc', {})
    if not toc_cfg.get('enabled', False):
        return

    try:
        toc_title = str(toc_cfg.get('title', 'ÍNDICE') or 'ÍNDICE')
        toc_depth = int(toc_cfg.get('depth', 2) or 2)
        toc_title_size = int(toc_cfg.get('title_size', 18) or 18)
        toc_title_color = str(toc_cfg.get('title_color', '#000000') or '#000000')
        toc_title_bold = bool(toc_cfg.get('title_bold', True))
        toc_title_italic = bool(toc_cfg.get('title_italic', False))
    except (ValueError, TypeError):
        toc_title = 'ÍNDICE'
        toc_depth = 2
        toc_title_size = 18
        toc_title_color = '#000000'
        toc_title_bold = True
        toc_title_italic = False

    body = doc.element.body
    font_name = config.get('font_name', 'Calibri')

    # Find the paragraph containing the first page break (end of cover)
    cover_break_p = None
    for p_elem in body.findall(qn('w:p')):
        for br in p_elem.findall('.//' + qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                cover_break_p = p_elem
                break
        if cover_break_p is not None:
            break

    if cover_break_p is None:
        print('TOC: No cover page break found, skipping')
        return

    # --- Build TOC elements ---
    toc_elements = []

    # Note: NO page break needed here — the cover page already ends with
    # a page break, so the TOC title starts on the next page directly.

    # 2. TOC Title paragraph
    p_title = OxmlElement('w:p')
    title_pPr = OxmlElement('w:pPr')
    title_jc = OxmlElement('w:jc')
    title_jc.set(qn('w:val'), 'center')
    title_pPr.append(title_jc)
    title_sp = OxmlElement('w:spacing')
    title_sp.set(qn('w:after'), '480')
    title_pPr.append(title_sp)
    p_title.append(title_pPr)

    title_r = OxmlElement('w:r')
    title_rPr = OxmlElement('w:rPr')
    if toc_title_bold:
        title_rPr.append(OxmlElement('w:b'))
    if toc_title_italic:
        title_rPr.append(OxmlElement('w:i'))
    _sz = OxmlElement('w:sz')
    _sz.set(qn('w:val'), str(toc_title_size * 2))
    title_rPr.append(_sz)
    _szCs = OxmlElement('w:szCs')
    _szCs.set(qn('w:val'), str(toc_title_size * 2))
    title_rPr.append(_szCs)
    _rFonts = OxmlElement('w:rFonts')
    _rFonts.set(qn('w:ascii'), font_name)
    _rFonts.set(qn('w:hAnsi'), font_name)
    title_rPr.append(_rFonts)
    try:
        rgb = hex_to_rgb(toc_title_color)
        col = OxmlElement('w:color')
        col.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*rgb))
        title_rPr.append(col)
    except Exception:
        pass
    title_r.append(title_rPr)
    title_t = OxmlElement('w:t')
    title_t.text = toc_title
    title_r.append(title_t)
    p_title.append(title_r)
    toc_elements.append(p_title)

    # 3. TOC Field paragraph (Word populates this from Heading styles)
    p_toc = OxmlElement('w:p')

    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1)
    p_toc.append(r1)

    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-{}" \\h \\z \\u '.format(toc_depth)
    r2.append(instr)
    p_toc.append(r2)

    r3 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'separate')
    r3.append(fc3)
    p_toc.append(r3)

    # Placeholder visible until Word updates the field
    r4 = OxmlElement('w:r')
    r4_rPr = OxmlElement('w:rPr')
    r4_rFonts = OxmlElement('w:rFonts')
    r4_rFonts.set(qn('w:ascii'), font_name)
    r4_rFonts.set(qn('w:hAnsi'), font_name)
    r4_rPr.append(r4_rFonts)
    r4.append(r4_rPr)
    t4 = OxmlElement('w:t')
    t4.text = 'Abra en Word para ver el índice completo'
    r4.append(t4)
    p_toc.append(r4)

    r5 = OxmlElement('w:r')
    fc5 = OxmlElement('w:fldChar')
    fc5.set(qn('w:fldCharType'), 'end')
    r5.append(fc5)
    p_toc.append(r5)
    toc_elements.append(p_toc)

    # 4. Page break after TOC (separates from body content)
    p_pb2 = OxmlElement('w:p')
    pb2_r = OxmlElement('w:r')
    pb2_br = OxmlElement('w:br')
    pb2_br.set(qn('w:type'), 'page')
    pb2_r.append(pb2_br)
    p_pb2.append(pb2_r)
    toc_elements.append(p_pb2)

    # --- Insert all TOC elements after cover break ---
    insert_point = cover_break_p
    for elem in toc_elements:
        insert_point.addnext(elem)
        insert_point = elem

    # --- Add updateFields setting so Word auto-updates TOC on open ---
    try:
        settings = doc.settings.element
        uf = settings.find(qn('w:updateFields'))
        if uf is None:
            uf = OxmlElement('w:updateFields')
            settings.append(uf)
        uf.set(qn('w:val'), 'true')
    except Exception:
        pass

    print('TOC: Inserted field code (depth={}, title="{}")'.format(toc_depth, toc_title))


def insert_cover_page(doc, pw=None, ph=None, config=None):
    """
    Insert a cover page at the start of the document.
    - Detect Title from first non-empty paragraph.
    - Insert uploaded cover image (if any).
    - Use different first page header/footer (no header/footer on cover).
    """
    if config is None:
        config = {}
    cover_cfg = config.get('cover', {})

    # 1. Detect Title
    title_text = "TÍTULO DEL DOCUMENTO"
    for p in doc.paragraphs:
        if p.text.strip():
            title_text = p.text.strip()
            break

    # 2. Check for Cover Image
    upload_folder = app.config['UPLOAD_FOLDER']
    cover_path = None
    for ext in ['.png', '.jpg', '.jpeg']:
        cand = os.path.join(upload_folder, 'custom_cover' + ext)
        if os.path.exists(cand):
            cover_path = cand
            break

    # 3. Enable Different First Page Header/Footer
    if doc.sections:
        section = doc.sections[0]
        if pw is None: pw = section.page_width or int(8.5 * 914400)
        if ph is None: ph = section.page_height or int(11 * 914400)
        tm = section.top_margin or Cm(2.5)
        section.different_first_page_header_footer = True
        
        # Title vertical position from config (default 55%)
        pos_y_pct = float(cover_cfg.get('pos_y', 55)) / 100.0
        title_target_y = ph * pos_y_pct
        title_space = max(0, title_target_y - tm)

    # 4. Read cover style settings
    cover_font = cover_cfg.get('font', 'Calibri')
    cover_size = int(cover_cfg.get('size', 36))
    cover_color_hex = cover_cfg.get('color', '#000000')
    cover_align_str = cover_cfg.get('align', 'center')
    cover_bold = cover_cfg.get('bold', True)
    cover_italic = cover_cfg.get('italic', False)

    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }
    cover_align = align_map.get(cover_align_str, WD_ALIGN_PARAGRAPH.CENTER)

    try:
        cover_rgb = hex_to_rgb(cover_color_hex)
    except Exception:
        cover_rgb = (0, 0, 0)

    # 5. Insert Content (Prepend)
    if not doc.paragraphs:
        doc.add_paragraph()
    
    p_first = doc.paragraphs[0]
    
    # Page Break (End of Cover)
    p_break = p_first.insert_paragraph_before()
    p_break.add_run().add_break(WD_BREAK.PAGE)
    
    # Title (uppercase)
    p_title = p_break.insert_paragraph_before()
    p_title.alignment = cover_align
    p_title.paragraph_format.line_spacing = 1.0
    p_title.paragraph_format.space_after = Pt(0)
    r = p_title.add_run(title_text.upper())
    r.font.name = cover_font

    # --- Auto font-size: reduce if title exceeds 3 lines at configured size ---
    cover_width_pct = float(cover_cfg.get('width', 80)) / 100.0
    title_width_emu = int((pw or int(8.5 * 914400)) * cover_width_pct)
    max_lines = 3
    min_font = 16  # never go below 16pt
    actual_size = cover_size

    def _estimate_lines(text, font_pt, width_emu):
        """Estimate how many lines a text would take given font size and width."""
        # Average character width ≈ 0.75 × font size (bold is wider; 1pt = 12700 EMU)
        char_width_emu = font_pt * 12700 * 0.75
        chars_per_line = max(1, int(width_emu / char_width_emu))
        # Word-wrap simulation
        words = text.split()
        lines = 1
        current_len = 0
        for word in words:
            word_len = len(word)
            if current_len == 0:
                current_len = word_len
            elif current_len + 1 + word_len <= chars_per_line:
                current_len += 1 + word_len
            else:
                lines += 1
                current_len = word_len
        return lines

    est_lines = _estimate_lines(title_text.upper(), actual_size, title_width_emu)
    while est_lines > max_lines and actual_size > min_font:
        actual_size -= 1
        est_lines = _estimate_lines(title_text, actual_size, title_width_emu)

    if actual_size != cover_size:
        print('COVER: auto-reduced font from {}pt to {}pt ({} lines)'.format(cover_size, actual_size, est_lines))

    r.font.size = Pt(actual_size)
    r.bold = cover_bold
    r.italic = cover_italic
    r.font.color.rgb = RGBColor(*cover_rgb)
    p_title.paragraph_format.space_before = Emu(int(title_space))
    print('COVER: pos_y={}%, space_before={} EMU ({:.1f} cm)'.format(
        int(pos_y_pct * 100), int(title_space), title_space / 360000))

    # Width and X-position: use left/right indent to constrain title box
    cover_width_pct = float(cover_cfg.get('width', 80)) / 100.0
    cover_pos_x_pct = float(cover_cfg.get('pos_x', 50)) / 100.0
    if pw:
        # Total page width available (without margins, but cover has no margins concept)
        available = pw
        title_width = int(available * cover_width_pct)
        # Calculate left indent so the title box is centered at pos_x
        title_left_edge = int(available * cover_pos_x_pct - title_width / 2)
        title_right_edge = int(available - title_left_edge - title_width)
        p_title.paragraph_format.left_indent = max(0, title_left_edge)
        p_title.paragraph_format.right_indent = max(0, title_right_edge)

    target = p_title
    
    # Cover Image (Background - Behind Text, Full Page)
    # Embed the anchor image directly into the title paragraph's run
    # to avoid an extra paragraph that shifts the title position.
    if cover_path:
        ext = os.path.splitext(cover_path)[1].lower()
        if ext in ['.png', '.jpg', '.jpeg']:
             _embed_raster_cover_full(p_title, cover_path, ext, width_emu=pw, height_emu=ph)


def insert_back_page(doc, pw=None, ph=None):
    """
    Append a final page with only a full-page background image.
    Uses a dedicated new section so it has no header, footer, or page numbers.
    """
    upload_folder = app.config['UPLOAD_FOLDER']
    back_path = None
    for ext in ['.png', '.jpg', '.jpeg']:
        cand = os.path.join(upload_folder, 'custom_backpage' + ext)
        if os.path.exists(cand):
            back_path = cand
            break

    if not back_path:
        return  # No back page image uploaded

    if doc.sections:
        section = doc.sections[0]
        if pw is None: pw = section.page_width or int(8.5 * 914400)
        if ph is None: ph = section.page_height or int(11 * 914400)

    # Add a new section for the back page (isolated: no header/footer)
    from docx.enum.section import WD_SECTION
    back_section = doc.add_section(WD_SECTION.NEW_PAGE)
    back_section.page_width = pw
    back_section.page_height = ph
    back_section.different_first_page_header_footer = True

    # Clear/unlink header and footer so nothing shows
    for hdr in [back_section.header, back_section.first_page_header]:
        hdr.is_linked_to_previous = False
        for p in hdr.paragraphs:
            for run in p.runs:
                run.text = ''

    for ftr in [back_section.footer, back_section.first_page_footer]:
        ftr.is_linked_to_previous = False
        for p in ftr.paragraphs:
            for run in p.runs:
                run.text = ''

    # Add the image paragraph in this new section
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(0)
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.line_spacing = Pt(0)

    ext = os.path.splitext(back_path)[1].lower()
    if ext in ['.png', '.jpg', '.jpeg']:
        _embed_raster_cover_full(p_img, back_path, ext, width_emu=pw, height_emu=ph)


def _embed_raster_cover_full(para, image_path, ext, width_emu, height_emu):
    """Embed PNG/JPG as a behind-text anchor image stretched to full page size."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    import lxml.etree as etree
    import uuid as _uuid

    content_type_map = {
        '.png': 'image/png', '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg', '.gif': 'image/gif', '.bmp': 'image/bmp',
    }
    content_type = content_type_map.get(ext.lower(), 'image/png')

    with open(image_path, 'rb') as f:
        image_bytes = f.read()

    footer_part = para.part
    uid = _uuid.uuid4().hex[:8]
    part_name = PackURI('/word/media/cover_{}{}'.format(uid, ext))
    image_part = Part(part_name, content_type, image_bytes)
    rId = footer_part.relate_to(
        image_part,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    )
    img_id = (abs(hash(uid)) % 65000) + 300

    # Anchor centered on page, full size
    anchor_xml = (
        '<w:drawing'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        ' relativeHeight="2" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        '<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="{img_id}" name="CoverImg{img_id}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic>'
        '<pic:nvPicPr>'
        '<pic:cNvPr id="{img_id}" name="CoverImg{img_id}"/>'
        '<pic:cNvPicPr/>'
        '</pic:nvPicPr>'
        '<pic:blipFill>'
        '<a:blip r:embed="{rId}"/>'
        '<a:stretch><a:fillRect/></a:stretch>'
        '</pic:blipFill>'
        '<pic:spPr>'
        '<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '</pic:spPr>'
        '</pic:pic>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:anchor>'
        '</w:drawing>'
    ).format(cx=int(width_emu), cy=int(height_emu), rId=rId, img_id=img_id)

    drawing_element = etree.fromstring(anchor_xml)
    p = para._p
    r_elem = OxmlElement('w:r')
    r_elem.append(drawing_element)
    p.append(r_elem)


def _add_fld_char(run_elem, fld_char_type):
    """Add a w:fldChar element to a run element."""
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), fld_char_type)
    run_elem.append(fld_char)


def _add_instr_text(run_elem, text):
    """Add a w:instrText element to a run element."""
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = text
    run_elem.append(instr_text)


def insert_page_numbers(doc, style='arabic', position='center', fmt='page_only', font_name='Calibri', font_size=11, toc_enabled=False):
    """
    Insert page numbers into the footer of body sections only.
    GLOBAL RULE: numbering starts at 1, excluding cover, TOC, and backpage.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para_align = align_map.get(position, WD_ALIGN_PARAGRAPH.CENTER)

    # Helper: create a run with font properties
    def make_run():
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size) * 2))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size) * 2))
        rPr.append(szCs)
        r.append(rPr)
        return r

    def add_page_field(para):
        """Add the page number field to a paragraph based on chosen format."""
        if fmt == 'page_only':
            r = make_run(); _add_fld_char(r, 'begin'); para._p.append(r)
            r2 = make_run(); _add_instr_text(r2, ' PAGE '); para._p.append(r2)
            r3 = make_run(); _add_fld_char(r3, 'end'); para._p.append(r3)
        elif fmt == 'page_of_total':
            r = make_run(); _add_fld_char(r, 'begin'); para._p.append(r)
            r2 = make_run(); _add_instr_text(r2, ' PAGE '); para._p.append(r2)
            r3 = make_run(); _add_fld_char(r3, 'end'); para._p.append(r3)
            rSep = make_run(); tSep = OxmlElement('w:t'); tSep.text = ' de '; tSep.set(qn('xml:space'), 'preserve'); rSep.append(tSep); para._p.append(rSep)
            r4 = make_run(); _add_fld_char(r4, 'begin'); para._p.append(r4)
            r5 = make_run(); _add_instr_text(r5, ' NUMPAGES '); para._p.append(r5)
            r6 = make_run(); _add_fld_char(r6, 'end'); para._p.append(r6)
        elif fmt == 'dash':
            rPre = make_run(); tPre = OxmlElement('w:t'); tPre.text = '- '; tPre.set(qn('xml:space'), 'preserve'); rPre.append(tPre); para._p.append(rPre)
            r = make_run(); _add_fld_char(r, 'begin'); para._p.append(r)
            r2 = make_run(); _add_instr_text(r2, ' PAGE '); para._p.append(r2)
            r3 = make_run(); _add_fld_char(r3, 'end'); para._p.append(r3)
            rPost = make_run(); tPost = OxmlElement('w:t'); tPost.text = ' -'; tPost.set(qn('xml:space'), 'preserve'); rPost.append(tPost); para._p.append(rPost)

    # --- Clean paragraph-level sectPr page restarts ---
    # Original documents may have section breaks (w:pPr/w:sectPr) with
    # pgNumType start values that restart page numbering at each break.
    # Remove these so numbering flows continuously through the document.
    body = doc.element.body
    for p_elem in body.findall(qn('w:p')):
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            continue
        p_sectPr = pPr.find(qn('w:sectPr'))
        if p_sectPr is None:
            continue
        pgNumType = p_sectPr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            if qn('w:start') in pgNumType.attrib:
                del pgNumType.attrib[qn('w:start')]

    num_sections = len(doc.sections)
    if num_sections == 0:
        return

    last_idx = num_sections - 1
    is_backpage = (doc.sections[last_idx].different_first_page_header_footer
                   if last_idx > 0 else False)

    first_body_done = False  # Track whether we've set up the first body section

    for i, section in enumerate(doc.sections):
        # Skip back page (last section with different_first_page)
        if i == last_idx and is_backpage:
            section.footer.is_linked_to_previous = False
            continue

        if i == 0:
            if toc_enabled:
                # With TOC: section 0 is ONLY the cover page, skip it
                continue
            else:
                # Without TOC: section 0 has cover + body.
                # Cover is hidden by different_first_page_header_footer.
                # Set start=0 so cover=page0(hidden), first body page=page1.
                sectPr = section._sectPr
                pgNumType = sectPr.find(qn('w:pgNumType'))
                if pgNumType is None:
                    pgNumType = OxmlElement('w:pgNumType')
                    sectPr.append(pgNumType)
                pgNumType.set(qn('w:start'), '0')
                footer = section.footer
                para = footer.add_paragraph()
                para.alignment = para_align
                add_page_field(para)
                first_body_done = True
                continue

        if i == 1 and toc_enabled:
            # With TOC: section 1 is the TOC page, skip it
            section.footer.is_linked_to_previous = False
            continue

        # Body section: add page numbers
        section.footer.is_linked_to_previous = False
        sectPr = section._sectPr

        if not first_body_done:
            # First body section: restart numbering from 1
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is None:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            pgNumType.set(qn('w:start'), '1')
            first_body_done = True
        else:
            # Subsequent body sections: remove any pgNumType start so numbering continues
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is not None:
                if qn('w:start') in pgNumType.attrib:
                    del pgNumType.attrib[qn('w:start')]

        footer = section.footer
        para = footer.add_paragraph()
        para.alignment = para_align
        add_page_field(para)

def _parse_emf_table(emf_bytes):
    """Parse EMF binary records to extract table data directly from vector drawing commands.
    Returns a dict {'rows': [...], 'bold_map': set()} or None if no table found.
    bold_map contains (row_idx, col_idx) tuples for cells that should be bold."""
    import struct as _struct

    EMR_EOF = 14
    EMR_MOVETOEX = 27
    EMR_SELECTOBJECT = 37
    EMR_LINETO = 54
    EMR_CREATEFONTINDIRECTW = 82
    EMR_EXTTEXTOUTW = 84

    # Verify EMF header
    if len(emf_bytes) < 88:
        return None
    rec_type = _struct.unpack_from('<I', emf_bytes, 0)[0]
    if rec_type != 1:  # EMR_HEADER
        return None
    sig = _struct.unpack_from('<I', emf_bytes, 40)[0]
    if sig != 0x464D4520:  # " EMF"
        return None

    texts = []          # (x, y, text, is_bold)
    h_segments = []     # (y, x_start, x_end)
    v_segments = []     # (x, y_start, y_end)
    last_move = None

    font_objects = {}
    current_bold = False
    offset = 0

    while offset < len(emf_bytes) - 8:
        rt = _struct.unpack_from('<I', emf_bytes, offset)[0]
        rs = _struct.unpack_from('<I', emf_bytes, offset + 4)[0]
        if rs < 8:
            break

        if rt == EMR_CREATEFONTINDIRECTW and rs >= 32:
            ih_object = _struct.unpack_from('<I', emf_bytes, offset + 8)[0]
            lf_weight = _struct.unpack_from('<i', emf_bytes, offset + 28)[0]
            font_objects[ih_object] = (lf_weight >= 700)

        elif rt == EMR_SELECTOBJECT and rs >= 12:
            ih_object = _struct.unpack_from('<I', emf_bytes, offset + 8)[0]
            if ih_object in font_objects:
                current_bold = font_objects[ih_object]

        elif rt == EMR_EXTTEXTOUTW and rs > 76:
            ref_x = _struct.unpack_from('<i', emf_bytes, offset + 36)[0]
            ref_y = _struct.unpack_from('<i', emf_bytes, offset + 40)[0]
            nchars = _struct.unpack_from('<I', emf_bytes, offset + 44)[0]
            off_string = _struct.unpack_from('<I', emf_bytes, offset + 48)[0]
            if nchars > 0 and off_string > 0:
                abs_off = offset + off_string
                if abs_off + nchars * 2 <= len(emf_bytes):
                    text = emf_bytes[abs_off:abs_off + nchars * 2].decode('utf-16-le', errors='replace')
                    if text.strip():
                        texts.append((ref_x, ref_y, text.strip(), current_bold))

        elif rt == EMR_MOVETOEX and rs >= 16:
            x = _struct.unpack_from('<i', emf_bytes, offset + 8)[0]
            y = _struct.unpack_from('<i', emf_bytes, offset + 12)[0]
            last_move = (x, y)

        elif rt == EMR_LINETO and rs >= 16 and last_move:
            x = _struct.unpack_from('<i', emf_bytes, offset + 8)[0]
            y = _struct.unpack_from('<i', emf_bytes, offset + 12)[0]
            x1, y1 = last_move
            if y1 == y:  # horizontal
                h_segments.append((y, min(x1, x), max(x1, x)))
            elif x1 == x:  # vertical
                v_segments.append((x, min(y1, y), max(y1, y)))
            last_move = (x, y)

        offset += rs
        if rt == EMR_EOF:
            break

    if not texts:
        return None

    # Need at least 2 horizontal lines and 2 vertical lines for a table grid
    unique_h = sorted(set(seg[0] for seg in h_segments))
    unique_v = sorted(set(seg[0] for seg in v_segments))

    if len(unique_h) < 2 or len(unique_v) < 2:
        y_vals = sorted(set(t[1] for t in texts))
        if len(y_vals) < 2:
            return None
        x_vals = sorted(set(t[0] for t in texts))
        if len(x_vals) < 2:
            return None

    # Group texts by row (Y coordinate)
    y_vals = sorted(set(t[1] for t in texts))

    # Determine columns from vertical lines or X positions
    if len(unique_v) >= 2:
        col_boundaries = unique_v
    else:
        all_x = sorted(set(t[0] for t in texts))
        col_boundaries = [all_x[0] - 1] + all_x + [all_x[-1] + 1]

    # Assign texts to columns
    num_cols = max(1, len(col_boundaries) - 1)
    rows_data = []
    bold_map = set()  # (row_idx, col_idx) for bold cells

    for row_idx, y in enumerate(y_vals):
        row_texts = [(t[0], t[2], t[3]) for t in texts if t[1] == y]
        row_texts.sort(key=lambda t: t[0])

        cells = [''] * num_cols
        for tx, txt, is_bold in row_texts:
            col_idx = num_cols - 1
            for ci in range(len(col_boundaries) - 1):
                if tx < col_boundaries[ci + 1]:
                    col_idx = ci
                    break
            if cells[col_idx]:
                cells[col_idx] += ' ' + txt
            else:
                cells[col_idx] = txt
            if is_bold:
                bold_map.add((row_idx, col_idx))
        rows_data.append(cells)

    if not rows_data or all(all(c == '' for c in row) for row in rows_data):
        return None

    # --- Detect merged cells from missing vertical borders ---
    tol = 3
    merge_map = {}  # (row_idx, col_idx) -> span count

    if len(unique_h) >= 2 and len(unique_v) >= 2:
        for row_idx, y_text in enumerate(y_vals):
            # Find horizontal lines above and below this text row
            y_above = None
            y_below = None
            for hy in unique_h:
                if hy <= y_text + tol:
                    y_above = hy
                elif y_below is None:
                    y_below = hy
            if y_above is None or y_below is None:
                continue

            # Check each interior column border
            for col_j in range(num_cols - 1):
                x_border = col_boundaries[col_j + 1]
                has_border = any(
                    abs(sx - x_border) <= tol and sy1 <= y_above + tol and sy2 >= y_below - tol
                    for sx, sy1, sy2 in v_segments
                )
                if not has_border:
                    # Find the merge start for this row
                    start = col_j
                    while start > 0 and (row_idx, start) not in merge_map and start - 1 >= 0:
                        # Check if previous col was already a merge start
                        if (row_idx, start - 1) in merge_map:
                            start = start - 1
                            break
                        start -= 1
                        break
                    if (row_idx, start) in merge_map:
                        merge_map[(row_idx, start)] += 1
                    else:
                        merge_map[(row_idx, col_j)] = 2

    # Consolidate text from merged cells into the start cell
    for (ri, ci), span in merge_map.items():
        for k in range(ci + 1, min(ci + span, num_cols)):
            if rows_data[ri][k]:
                rows_data[ri][ci] = (rows_data[ri][ci] + ' ' + rows_data[ri][k]).strip()
                rows_data[ri][k] = ''

    print('EMF: Parsed {} rows x {} cols, {} merges from vector data'.format(
        len(rows_data), num_cols, len(merge_map)))
    return {
        'rows': rows_data,
        'bold_map': bold_map,
        'merge_map': merge_map,
        'col_boundaries': col_boundaries
    }


def ocr_extract_tables(doc):
    """Scan images in doc, run OCR table detection, replace with native tables.
    Handles EMF images by parsing vector data directly (no OCR needed).
    Handles raster images (PNG/JPG) via Tesseract OCR."""
    shapes_to_process = []

    # Collect inline shapes (images)
    for para in doc.paragraphs:
        for run in para.runs:
            if run._r.findall(qn('w:drawing')):
                shapes_to_process.append((para, run))

    print('OCR: Found {} image(s) to scan for tables'.format(len(shapes_to_process)))

    replaced = 0
    for para, run in shapes_to_process:
        try:
            # Try to find the image relationship
            blip_list = run._r.findall('.//' + qn('a:blip'))
            if not blip_list:
                continue
            rId = blip_list[0].get(qn('r:embed'))
            if not rId:
                continue

            image_part = doc.part.related_parts.get(rId)
            if not image_part:
                continue

            img_bytes = image_part.blob
            content_type = getattr(image_part, 'content_type', 'unknown')
            print('OCR: Processing image ({}, {} bytes)'.format(content_type, len(img_bytes)))

            rows_data = None
            bold_map = set()
            merge_map = {}
            col_boundaries = None

            # --- EMF/WMF: parse vector data directly (no OCR needed) ---
            if content_type and ('emf' in content_type.lower() or 'wmf' in content_type.lower()):
                print('OCR: EMF/WMF detected — attempting direct vector parsing')
                result = _parse_emf_table(img_bytes)
                if result is None:
                    print('OCR: No table structure found in EMF')
                    continue
                rows_data = result['rows']
                bold_map = result['bold_map']
                merge_map = result.get('merge_map', {})
                col_boundaries = result.get('col_boundaries')

            # --- Raster (PNG/JPG): use Tesseract OCR ---
            else:
                if not HAS_OCR:
                    print("OCR: Skipping raster image (img2table/pytesseract not installed)")
                    continue

                from PIL import Image as PILImage
                try:
                    pil_img = PILImage.open(io.BytesIO(img_bytes))
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                        pil_img.save(tmp, format='PNG')
                        tmp_path = tmp.name
                except Exception as e:
                    print('OCR: Cannot decode image: {}'.format(str(e)))
                    continue

                try:
                    ocr = TesseractOCR(n_threads=1, lang='spa+eng')
                    img_doc = Img2TableImage(src=tmp_path)
                    tables = img_doc.extract_tables(ocr=ocr, implicit_rows=True, implicit_columns=True)

                    # Fallback: try borderless detection
                    if not tables:
                        print('OCR: No bordered table found, trying borderless detection...')
                        tables = img_doc.extract_tables(ocr=ocr, implicit_rows=True, implicit_columns=True, borderless_tables=True)

                    if not tables:
                        print('OCR: WARNING — No table detected in PNG image ({} bytes), skipping'.format(len(img_bytes)))
                        continue

                    table_data = tables[0].df
                    if table_data is None or table_data.empty:
                        continue

                    rows_data = [table_data.columns.tolist()] + table_data.values.tolist()
                finally:
                    os.unlink(tmp_path)

            # --- Build native DOCX table ---
            if rows_data is None:
                continue

            num_cols = len(rows_data[0]) if rows_data else 0
            num_rows = len(rows_data)
            if num_cols == 0 or num_rows == 0:
                continue

            # Read original image width from the drawing XML (EMU units)
            img_width_emu = None
            try:
                drawings = run._r.findall('.//' + qn('wp:inline'))
                if not drawings:
                    drawings = run._r.findall('.//' + qn('wp:anchor'))
                if drawings:
                    extent = drawings[0].find(qn('wp:extent'))
                    if extent is not None:
                        img_width_emu = int(extent.get('cx', '0'))
            except Exception:
                pass

            tbl = doc.add_table(rows=num_rows, cols=num_cols)
            tbl.autofit = True

            # Set table width to match original image
            tbl_pr = tbl._tbl.tblPr
            tbl_width = tbl_pr.find(qn('w:tblW'))
            if tbl_width is None:
                tbl_width = OxmlElement('w:tblW')
                tbl_pr.append(tbl_width)

            if img_width_emu and img_width_emu > 0:
                width_twips = int(img_width_emu / 635)
                tbl_width.set(qn('w:w'), str(width_twips))
                tbl_width.set(qn('w:type'), 'dxa')

                # Set proportional column widths from EMF grid boundaries
                if col_boundaries and len(col_boundaries) == num_cols + 1:
                    total_grid = col_boundaries[-1] - col_boundaries[0]
                    if total_grid > 0:
                        for ci in range(num_cols):
                            col_w = col_boundaries[ci + 1] - col_boundaries[ci]
                            col_twips = int(width_twips * col_w / total_grid)
                            for ri in range(num_rows):
                                tc = tbl.cell(ri, ci)._tc
                                tcPr = tc.get_or_add_tcPr()
                                tcW = tcPr.find(qn('w:tcW'))
                                if tcW is None:
                                    tcW = OxmlElement('w:tcW')
                                    tcPr.append(tcW)
                                tcW.set(qn('w:w'), str(col_twips))
                                tcW.set(qn('w:type'), 'dxa')
            else:
                tbl_width.set(qn('w:w'), '0')
                tbl_width.set(qn('w:type'), 'auto')

            # Fill cell text and bold
            for i, row in enumerate(rows_data):
                for j, cell_val in enumerate(row):
                    cell = tbl.cell(i, j)
                    cell_text = str(cell_val) if cell_val is not None else ''
                    cell.text = cell_text
                    if (i, j) in bold_map:
                        for r in cell.paragraphs[0].runs:
                            r.font.bold = True

            # Apply smart merges detected from EMF grid (missing vertical borders)
            for (ri, ci), span in merge_map.items():
                end_col = min(ci + span - 1, num_cols - 1)
                if end_col > ci:
                    try:
                        tbl.cell(ri, ci).merge(tbl.cell(ri, end_col))
                    except Exception:
                        pass

            # Move table to replace the image paragraph
            para._p.addnext(tbl._tbl)

            # Remove the original image paragraph
            parent = para._p.getparent()
            if parent is not None:
                parent.remove(para._p)

            replaced += 1

        except Exception as e:
            print('OCR error on image: {}'.format(str(e)))
            continue

    print('OCR: replaced {} image(s) with native tables'.format(replaced))


def _fix_hyphenated_words(doc):
    """Rejoin words split by line-break hyphens (e.g. 'cons- trucción' → 'construcción').
    Common in PDF-to-DOCX conversions where syllable breaks become literal text.
    Handles multiple patterns across runs and within single runs."""
    import re as _re_hyph
    _LETTER = r'[a-záéíóúñüàèìòùâêîôûäëïöü]'
    _SPLIT_RE = _re_hyph.compile(
        r'(' + _LETTER + r'+)-\s+(' + _LETTER + r'+)', _re_hyph.IGNORECASE)
    fixed = 0

    for paragraph in doc.paragraphs:
        # --- Pass 1: fix splits within single runs ---
        for run in paragraph.runs:
            new_text, count = _SPLIT_RE.subn(r'\1\2', run.text)
            if count:
                run.text = new_text
                fixed += count

        # --- Pass 2: fix splits across adjacent runs ---
        runs = list(paragraph.runs)
        if len(runs) < 2:
            continue
        i = 0
        while i < len(runs) - 1:
            r_text = runs[i].text
            next_text = runs[i + 1].text if i + 1 < len(runs) else ''

            # Pattern A: current run ends with "word- " → next run starts lowercase
            m = _re_hyph.search(r'(' + _LETTER + r'+)-\s*$', r_text, _re_hyph.IGNORECASE)
            if m and next_text and next_text[0].islower():
                runs[i].text = r_text[:m.start()] + m.group(1) + next_text
                paragraph._p.remove(runs[i + 1]._r)
                runs.pop(i + 1)
                fixed += 1
                continue

            # Pattern B: current run is just "- " → merge prev + next
            if r_text.strip() == '-' and i > 0 and next_text:
                prev_text = runs[i - 1].text
                if (prev_text and prev_text[-1].isalpha() and
                    next_text and next_text[0].islower()):
                    runs[i - 1].text = prev_text + next_text
                    paragraph._p.remove(runs[i]._r)
                    paragraph._p.remove(runs[i + 1]._r)
                    runs.pop(i + 1)
                    runs.pop(i)
                    fixed += 1
                    continue

            # Pattern C: next run starts with "- word" → merge into current
            m2 = _re_hyph.match(r'^-\s*(' + _LETTER + r'+)', next_text, _re_hyph.IGNORECASE)
            if m2 and r_text and r_text[-1].isalpha():
                remainder = next_text[m2.end():]
                runs[i].text = r_text + m2.group(1)
                if remainder.strip():
                    runs[i + 1].text = remainder
                else:
                    paragraph._p.remove(runs[i + 1]._r)
                    runs.pop(i + 1)
                fixed += 1
                continue

            i += 1

    if fixed:
        print('DEHYPHEN: Fixed {} split words'.format(fixed))


def apply_styles(doc, config, paper_size='letter'):

    # --- Paper Size ---
    paper_sizes = {
        'letter': (Inches(8.5), Inches(11)),
        'a4': (Cm(21), Cm(29.7)),
        'legal': (Inches(8.5), Inches(14)),
    }
    pw, ph = paper_sizes.get(paper_size, paper_sizes['letter'])

    # Fix hyphenated/split words before any other processing
    _fix_hyphenated_words(doc)

    # Convert internal nextPage section breaks to continuous to prevent blank pages.
    body = doc.element.body
    # Paragraph-level sectPr
    for p_elem in body.findall(qn('w:p')):
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            continue
        p_sectPr = pPr.find(qn('w:sectPr'))
        if p_sectPr is None:
            continue
        type_el = p_sectPr.find(qn('w:type'))
        if type_el is None:
            type_el = OxmlElement('w:type')
            p_sectPr.append(type_el)
        type_el.set(qn('w:val'), 'continuous')
    # Body-level sectPr (last section)
    body_sectPr = body.find(qn('w:sectPr'))
    if body_sectPr is not None:
        type_el = body_sectPr.find(qn('w:type'))
        if type_el is None:
            type_el = OxmlElement('w:type')
            body_sectPr.append(type_el)
        if type_el.get(qn('w:val')) != 'continuous':
            type_el.set(qn('w:val'), 'continuous')
    # Clean stale header content and flags from original sections (PDF-to-DOCX artifacts).
    for section in doc.sections:
        if section.different_first_page_header_footer:
            section.different_first_page_header_footer = False
        header = section.header
        if header.is_linked_to_previous:
            continue
        has_text = any(p.text.strip() for p in header.paragraphs)
        old_imgs = header._element.findall('.//' + qn('w:drawing')) + header._element.findall('.//' + qn('w:pict'))
        if has_text or old_imgs:
            for para in header.paragraphs:
                p = para._p
                for child in list(p):
                    p.remove(child)
        try:
            fph = section.first_page_header
            for para in fph.paragraphs:
                p = para._p
                for child in list(p):
                    p.remove(child)
        except Exception:
            pass

    # --- Word Compatibility Mode (Word 2013+, better justify spacing) ---
    try:
        settings = doc.settings.element
        
        # Disable evenAndOddHeaders globally so pages 2, 4, 6... don't fall back to an empty/broken header
        even_odd = settings.find(qn('w:evenAndOddHeaders'))
        if even_odd is not None:
            settings.remove(even_odd)

        compat = settings.find(qn('w:compat'))
        if compat is None:
            compat = OxmlElement('w:compat')
            settings.append(compat)
        # Remove old compatibilityMode if present
        for cs in compat.findall(qn('w:compatSetting')):
            if cs.get(qn('w:name')) == 'compatibilityMode':
                compat.remove(cs)
        cs = OxmlElement('w:compatSetting')
        cs.set(qn('w:name'), 'compatibilityMode')
        cs.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
        cs.set(qn('w:val'), '15')
        compat.append(cs)
    except Exception:
        pass

    # --- Cover Page ---
    insert_cover_page(doc, pw, ph, config=config)
    insert_back_page(doc, pw, ph)

    # --- Table of Contents (if enabled) ---
    insert_toc_page(doc, config)

    # --- OCR Table Extraction (if enabled) ---
    table_cfg = config.get('tables', {})
    if table_cfg.get('ocr_tables', False):
        try:
            ocr_extract_tables(doc)
        except Exception as e:
            print('OCR table extraction failed: {}'.format(str(e)))

    for section in doc.sections:
        section.page_width = pw
        section.page_height = ph
        # Dynamic margins from config (defaults to Unifranz standard)
        margins = config.get('margins', {})
        section.top_margin    = Cm(float(margins.get('top',    3.7)))
        section.bottom_margin = Cm(float(margins.get('bottom', 3.0)))
        section.left_margin   = Cm(float(margins.get('left',   3.0)))
        section.right_margin  = Cm(float(margins.get('right',  3.0)))

    # --- Header & Footer Images + Page Numbers ---
    # Page numbers first so they appear ABOVE the footer image
    clean_footers(doc)
    embed_header_image(doc)

    # --- Page Numbers (inserted before footer image) ---
    pn_cfg = config.get('page_numbers', {})
    font_name = config.get('font_name', 'Calibri')
    font_size = int(config.get('font_size', 11))
    toc_enabled = config.get('toc', {}).get('enabled', False)
    if pn_cfg.get('enabled', True):
        insert_page_numbers(doc,
            style    = pn_cfg.get('style',    'arabic'),
            position = pn_cfg.get('position', 'center'),
            fmt      = pn_cfg.get('format',   'page_only'),
            font_name = font_name,
            font_size = font_size,
            toc_enabled = toc_enabled)

    # Footer image goes AFTER page numbers
    embed_footer_image(doc)

    font_name = config.get('font_name', 'Arial')
    font_size = int(config.get('font_size', 11))
    line_spacing_val = float(config.get('line_spacing', 1.15))
    text_align_str = config.get('text_align', 'left')

    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'distribute': WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    text_align = align_map.get(text_align_str, WD_ALIGN_PARAGRAPH.LEFT)

    # Apply to Normal style
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = font_name
        normal_style.font.size = Pt(font_size)
        # Set document language to Spanish
        rPr = normal_style.element.get_or_add_rPr()
        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = OxmlElement('w:lang')
            rPr.append(lang)
        lang.set(qn('w:val'), 'es-BO')
        lang.set(qn('w:eastAsia'), 'es-BO')
        lang.set(qn('w:bidi'), 'es-BO')
    except Exception:
        pass

    link_color_hex = config.get('link_color', '#0563C1')
    try:
        link_rgb = hex_to_rgb(link_color_hex)
    except Exception:
        link_rgb = (5, 99, 193)  # default blue

    # --- Paragraphs (body text) ---
    def _apply_body_paragraph(paragraph):
        """Apply body font/size/align to a non-heading paragraph."""
        is_heading = paragraph.style.name.startswith('Heading')

        # Check if paragraph is a standalone image (has drawing, no text)
        has_drawing = bool(paragraph._p.findall(qn('w:r') + '/' + qn('w:drawing')))
        if not has_drawing:
            has_drawing = bool(paragraph._p.findall('.//' + qn('w:drawing')))
        has_text = any(r.text.strip() for r in paragraph.runs)

        if has_drawing and not has_text:
            # Standalone image — center it
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif not is_heading:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
            paragraph.alignment = text_align

        # Line spacing on all paragraphs
        pPr = paragraph._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:line'), str(int(line_spacing_val * 240)))
        spacing.set(qn('w:lineRule'), 'auto')

        # Set paragraph-level rPr font so Word auto-numbering (list numbers) inherits it
        p_rPr = pPr.find(qn('w:rPr'))
        if p_rPr is None:
            p_rPr = OxmlElement('w:rPr')
            pPr.insert(0, p_rPr)
        p_rFonts = p_rPr.find(qn('w:rFonts'))
        if p_rFonts is None:
            p_rFonts = OxmlElement('w:rFonts')
            p_rPr.append(p_rFonts)
        p_rFonts.set(qn('w:ascii'), font_name)
        p_rFonts.set(qn('w:hAnsi'), font_name)
        p_rFonts.set(qn('w:cs'), font_name)
        p_sz = p_rPr.find(qn('w:sz'))
        if p_sz is None:
            p_sz = OxmlElement('w:sz')
            p_rPr.append(p_sz)
        p_sz.set(qn('w:val'), str(font_size * 2))
        p_szCs = p_rPr.find(qn('w:szCs'))
        if p_szCs is None:
            p_szCs = OxmlElement('w:szCs')
            p_rPr.append(p_szCs)
        p_szCs.set(qn('w:val'), str(font_size * 2))

        # Widow/Orphan control: avoid single line at top/bottom of page
        wc = pPr.find(qn('w:widowControl'))
        if wc is None:
            wc = OxmlElement('w:widowControl')
            pPr.append(wc)
        wc.set(qn('w:val'), '1')

    # Skip cover page paragraphs (before first page break)
    past_cover_body = False
    for paragraph in doc.paragraphs:
        if not past_cover_body:
            for run in paragraph.runs:
                for br in run._r.findall(qn('w:br')):
                    if br.get(qn('w:type')) == 'page':
                        past_cover_body = True
            continue
        _apply_body_paragraph(paragraph)

    # Force configured font on numbering definitions (abstractNum levels)
    # Skip levels that use symbol fonts (bullets like •, □, ■ live in these)
    SYMBOL_FONTS = {'symbol', 'wingdings', 'wingdings 2', 'wingdings 3', 'webdings', 'courier new'}
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is not None:
            numbering_elm = numbering_part._element
            for lvl in numbering_elm.findall('.//' + qn('w:lvl')):
                rPr = lvl.find(qn('w:rPr'))
                # Check if this level currently uses a symbol font
                if rPr is not None:
                    existing_fonts = rPr.find(qn('w:rFonts'))
                    if existing_fonts is not None:
                        cur_font = (existing_fonts.get(qn('w:ascii')) or '').lower()
                        if cur_font in SYMBOL_FONTS:
                            continue  # Don't touch symbol bullet fonts

                # Check numFmt — only override 'decimal' (1,2,3) formats, not 'bullet'
                numFmt = lvl.find(qn('w:numFmt'))
                if numFmt is not None and numFmt.get(qn('w:val')) == 'bullet':
                    continue  # Skip bullet levels entirely

                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    lvl.append(rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:cs'), font_name)
    except Exception:
        pass  # No numbering part in this document

    # Apply font/size/spacing to paragraphs inside table cells too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _apply_body_paragraph(para)
                    # Zero paragraph spacing inside table cells to avoid extra padding
                    pPr = para._p.get_or_add_pPr()
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    spacing.set(qn('w:before'), '0')
                    spacing.set(qn('w:after'), '0')

    # Apply link color to all hyperlinks in body paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Check if this run is inside a hyperlink (w:hyperlink)
            parent = run._r.getparent()
            if parent is not None and parent.tag.endswith('}hyperlink'):
                run.font.color.rgb = RGBColor(*link_rgb)
                run.font.underline = True

    # --- Headings ---
    headings_cfg = config.get('headings', {})
    heading_style_map = {'Heading 1': 'h1', 'Heading 2': 'h2', 'Heading 3': 'h3'}

    # Ensure Heading 1/2/3 styles exist in the document's style gallery.
    # Many source documents only have 'Normal' — if the styles are missing,
    # all heading assignments fail silently (KeyError caught).
    # CRITICAL: We must also set w:outlineLvl so Word's TOC field (\o switch)
    # can find them — custom styles lack outline levels by default.
    import copy as _copy
    from docx.enum.style import WD_STYLE_TYPE
    _heading_outline = {'Heading 1': '0', 'Heading 2': '1', 'Heading 3': '2'}
    _heading_styles_available = {}
    for heading_name, outline_lvl in _heading_outline.items():
        try:
            h_style = doc.styles[heading_name]
            _heading_styles_available[heading_name] = True
        except KeyError:
            # Don't create new styles — the document may use latent/built-in
            # heading styles not present in styles.xml. Creating a duplicate
            # style corrupts numbering definitions (numId references scramble).
            _heading_styles_available[heading_name] = False
            continue
        pPr = h_style.element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            h_style.element.append(pPr)
        # Remove numPr from style definition to prevent style-level numbering
        # from overriding paragraph-level numbering
        for numPr in pPr.findall(qn('w:numPr')):
            pPr.remove(numPr)
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is None:
            outlineLvl = OxmlElement('w:outlineLvl')
            pPr.append(outlineLvl)
        outlineLvl.set(qn('w:val'), outline_lvl)

    import re as _re

    def _is_all_bold(paragraph):
        """Returns True if paragraph has runs and all visible runs are bold."""
        runs = [r for r in paragraph.runs if r.text.strip()]
        return bool(runs) and all(r.font.bold for r in runs)

    def _detect_heading_level(paragraph, is_first=False):
        """
        Returns (level, is_numbered).
        level: 'h1', 'h2', 'h3', or None
        is_numbered: True → numbered pattern or existing Heading style (appears in TOC)
                     False → bold-short-text fallback or doc title (styling only, NOT in TOC)
        """
        if is_first:
            return ('h1', False)  # doc_title: styled but NOT in TOC
        sname = paragraph.style.name
        if sname == 'Heading 1':
            return ('h1', True)
        if sname == 'Heading 2':
            return ('h2', True)
        if sname == 'Heading 3':
            return ('h3', True)
        if sname.startswith('Heading'):
            return ('h2', True)
        # Heuristic for documents without heading styles
        text = paragraph.text.strip()
        if not text or not _is_all_bold(paragraph):
            return (None, False)
        # Multi-level numbering (check deeper levels FIRST)
        if _re.match(r'^\d+\.\d+\.\d+[\.)\s]', text):
            return ('h3', True)
        if _re.match(r'^\d+\.\d+[\.)\s]', text):
            return ('h2', True)
        if _re.match(r'^\d+[\.\)]\s+\S', text):
            return ('h1', True)
        # Sub-heading: bold short paragraph — styling only, NOT in TOC
        if len(text) < 120 and not text.startswith('-') and not text.startswith('*'):
            return ('h2', False)
        return (None, False)

    def _apply_heading_cfg(paragraph, h_cfg, override_align=None):
        """Apply title/heading config dict to all runs of a paragraph."""
        for run in paragraph.runs:
            font_override = h_cfg.get('font', '')
            run.font.name = font_override if font_override else font_name
            size_val = h_cfg.get('size', '')
            if size_val and str(size_val).strip():
                try:
                    run.font.size = Pt(int(str(size_val).strip()))
                except (ValueError, TypeError):
                    pass
            color_val = h_cfg.get('color', '')
            if color_val and str(color_val).strip():
                try:
                    rgb = hex_to_rgb(str(color_val).strip())
                    run.font.color.rgb = RGBColor(*rgb)
                    # Clear theme color that may override explicit RGB
                    rPr = run._r.get_or_add_rPr()
                    color_elem = rPr.find(qn('w:color'))
                    if color_elem is not None:
                        for attr in ['w:themeColor', 'w:themeShade', 'w:themeTint']:
                            key = qn(attr)
                            if key in color_elem.attrib:
                                del color_elem.attrib[key]
                except Exception:
                    pass
            bold_val = h_cfg.get('bold')
            if bold_val is not None:
                run.font.bold = bool(bold_val)
            italic_val = h_cfg.get('italic')
            if italic_val is not None:
                run.font.italic = bool(italic_val)
        # Paragraph-level alignment for doc_title
        if override_align:
            align_map_local = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            paragraph.alignment = align_map_local.get(override_align, WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.page_break_before = False
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        # Set paragraph-level rPr color so Word auto-numbering (list numbers) inherits it
        color_val = h_cfg.get('color', '')
        if color_val and str(color_val).strip():
            try:
                rgb = hex_to_rgb(str(color_val).strip())
                pPr = paragraph._p.get_or_add_pPr()
                rPr = pPr.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    pPr.insert(0, rPr)
                color_elem = rPr.find(qn('w:color'))
                if color_elem is None:
                    color_elem = OxmlElement('w:color')
                    rPr.append(color_elem)
                color_elem.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*rgb))
            except Exception:
                pass

    # Apply config to heading styles in the style gallery (for future-proofing)
    for style_name, key in heading_style_map.items():
        h_cfg = headings_cfg.get(key, {})
        if not h_cfg:
            continue
        try:
            h_style = doc.styles[style_name]
            h_font = h_style.font
            h_font.name = font_name
            size_val = h_cfg.get('size', '')
            if size_val and str(size_val).strip():
                h_font.size = Pt(int(str(size_val).strip()))
            color_val = h_cfg.get('color', '')
            if color_val and str(color_val).strip():
                rgb = hex_to_rgb(str(color_val).strip())
                h_font.color.rgb = RGBColor(*rgb)
            bold_val = h_cfg.get('bold')
            if bold_val is not None:
                h_font.bold = bool(bold_val)
        except Exception:
            pass

    # Apply to every paragraph (style-based + heuristic, first paragraph = doc_title)
    # Skip cover page paragraphs (before the first page break) since they are
    # styled by insert_cover_page() and must not be overwritten.
    # Also skip TOC-injected paragraphs (section breaks + TOC field content)
    # which sit between the cover page break and the actual body content.
    first_found = False
    past_cover = False
    past_toc = False
    toc_enabled = config.get('toc', {}).get('enabled', False)
    for paragraph in doc.paragraphs:
        # Detect end of cover page (page break paragraph)
        if not past_cover:
            has_page_break = False
            for run in paragraph.runs:
                for br in run._r.findall(qn('w:br')):
                    if br.get(qn('w:type')) == 'page':
                        has_page_break = True
            if has_page_break:
                past_cover = True
            continue  # Skip all cover page paragraphs

        # Skip TOC-injected paragraphs (page break, title, field, page break)
        if toc_enabled and not past_toc:
            # Check for TOC field (fldChar)
            has_fld = bool(paragraph._p.findall('.//' + qn('w:fldChar')))
            is_toc_title = paragraph.text.strip() == config.get('toc', {}).get('title', 'ÍNDICE')
            if has_fld or is_toc_title:
                continue  # Skip TOC elements
            # Empty paragraph or page-break-only paragraph in TOC zone
            text = paragraph.text.strip()
            if not text:
                # Check if this is a page-break paragraph (part of TOC)
                has_pb = bool(paragraph._p.findall('.//' + qn('w:br')))
                if has_pb:
                    continue
            else:
                past_toc = True  # First real content paragraph after TOC

        is_first = False
        if not first_found and paragraph.text.strip():
            is_first = True
            first_found = True

        level, is_numbered = _detect_heading_level(paragraph, is_first=is_first)
        if not level:
            continue

        if is_first:
            # Apply doc_title config separately from h1
            dt_cfg = headings_cfg.get('doc_title', {})
            if dt_cfg:
                _apply_heading_cfg(paragraph, dt_cfg, override_align=dt_cfg.get('align', ''))
            # Doc title does NOT get Heading style — stays out of TOC
        else:
            # Only assign Word Heading style for NUMBERED headings (they appear in TOC)
            if is_numbered:
                style_name = {'h1': 'Heading 1', 'h2': 'Heading 2', 'h3': 'Heading 3'}.get(level, 'Heading 2')
                current_sname = paragraph.style.name if paragraph.style else ''
                is_already_heading = (current_sname.startswith('Heading') or
                                      current_sname.startswith('Título') or
                                      current_sname.startswith('Titulo'))
                if is_already_heading:
                    # Paragraph already has a heading style — do NOT reassign
                    # to avoid corrupting numbering definitions and style refs.
                    # Set outlineLvl at paragraph level to ensure TOC inclusion.
                    pPr = paragraph._p.get_or_add_pPr()
                    outlineLvl = pPr.find(qn('w:outlineLvl'))
                    if outlineLvl is None:
                        outlineLvl = OxmlElement('w:outlineLvl')
                        pPr.append(outlineLvl)
                    outlineLvl.set(qn('w:val'), {'h1': '0', 'h2': '1', 'h3': '2'}.get(level, '1'))
                elif _heading_styles_available.get(style_name, False):
                    try:
                        # Save existing numPr before style change
                        pPr = paragraph._p.find(qn('w:pPr'))
                        saved_numPr = None
                        if pPr is not None:
                            numPr_el = pPr.find(qn('w:numPr'))
                            if numPr_el is not None:
                                saved_numPr = _copy.deepcopy(numPr_el)
                        paragraph.style = doc.styles[style_name]
                        # Restore numPr if it existed
                        if saved_numPr is not None:
                            pPr = paragraph._p.get_or_add_pPr()
                            for np in pPr.findall(qn('w:numPr')):
                                pPr.remove(np)
                            pPr.append(saved_numPr)
                    except KeyError:
                        pass
                else:
                    # Style not available — set outlineLvl directly on paragraph
                    pPr = paragraph._p.get_or_add_pPr()
                    outlineLvl = pPr.find(qn('w:outlineLvl'))
                    if outlineLvl is None:
                        outlineLvl = OxmlElement('w:outlineLvl')
                        pPr.append(outlineLvl)
                    outlineLvl.set(qn('w:val'), {'h1': '0', 'h2': '1', 'h3': '2'}.get(level, '1'))
            # Apply heading formatting (font/size/color/bold) regardless
            if headings_cfg.get(level):
                _apply_heading_cfg(paragraph, headings_cfg[level])

    # --- Tables ---
    tables_cfg = config.get('tables', {})
    if tables_cfg and doc.tables:
        header_bg = tables_cfg.get('header_bg', '#4E2A84')
        header_text_color = tables_cfg.get('header_text', '#FFFFFF')
        border_v = border_style_to_val(tables_cfg.get('border_v', 'single'))
        border_h = border_style_to_val(tables_cfg.get('border_h', 'single'))
        # Separate colors per border type
        border_v_color_hex = tables_cfg.get('border_v_color', '#000000').lstrip('#')
        border_h_color_hex = tables_cfg.get('border_h_color', '#000000').lstrip('#')
        border_outline_color_hex = tables_cfg.get('border_outline_color', '#000000').lstrip('#')
        # Outline thickness: sz in eighths of a point (e.g. 4 = 0.5pt, 8 = 1pt, 16 = 2pt)
        border_outline_sz = str(tables_cfg.get('border_outline_sz', 4))
        zebra = tables_cfg.get('zebra', True)
        zebra_color_hex = tables_cfg.get('zebra_color', '#f1f5f9').lstrip('#').upper()
        align_numbers = tables_cfg.get('align_numbers', True)

        header_bg_rgb = hex_to_rgb(header_bg)
        header_text_rgb = hex_to_rgb(header_text_color)

        # --- Process Tables ---
        # Iterate through all tables to apply formatting
        for table in doc.tables:
            table.autofit = True
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Use first row as header if available
            header_cells = table.rows[0].cells
            for row_idx, row in enumerate(table.rows):
                is_header = (row_idx == 0)
                is_zebra_row = zebra and (row_idx % 2 == 0) and not is_header

                # Prevent row splitting across pages
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                cantSplit = trPr.find(qn('w:cantSplit'))
                if cantSplit is None:
                    cantSplit = OxmlElement('w:cantSplit')
                    trPr.append(cantSplit)

                num_rows = len(table.rows)
                
                for cell_idx, cell in enumerate(row.cells):
                    num_cols = len(row.cells)
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()

                    # Remove any existing noWrap — only add it for number cells below
                    existing_noWrap = tcPr.find(qn('w:noWrap'))
                    if existing_noWrap is not None:
                        tcPr.remove(existing_noWrap)

                    # Background color
                    if is_header or is_zebra_row:
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        if is_header:
                            fill_color = '{:02X}{:02X}{:02X}'.format(*header_bg_rgb)
                        else:
                            fill_color = zebra_color_hex
                        shd.set(qn('w:fill'), fill_color)
                        tcPr.append(shd)

                    # Cell text
                    for para in cell.paragraphs:
                        # Clear any paragraph-level shading (white highlight)
                        pPr = para._p.find(qn('w:pPr'))
                        if pPr is not None:
                            for shd_elem in pPr.findall(qn('w:shd')):
                                pPr.remove(shd_elem)
                        for run in para.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            run.font.highlight_color = None
                            # Clear run-level shading
                            rPr = run._r.find(qn('w:rPr'))
                            if rPr is not None:
                                for shd_elem in rPr.findall(qn('w:shd')):
                                    rPr.remove(shd_elem)
                            if is_header:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(*header_text_rgb)

                        # Remove all line spacing from table cells
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.line_spacing = 1.0

                        # Number alignment + no-wrap (prevent mid-number line breaks)
                        if align_numbers and not is_header:
                            text_stripped = cell.text.strip().replace(',', '.').replace('%', '').replace(' ', '')
                            try:
                                float(text_stripped)
                                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                # Also prevent wrapping for number cells
                                _noWrap = tcPr.find(qn('w:noWrap'))
                                if _noWrap is None:
                                    _noWrap = OxmlElement('w:noWrap')
                                    tcPr.append(_noWrap)
                            except ValueError:
                                pass

                        # Prevent any wrapping inside table cells (especially numbers)
                        for run in para.runs:
                            if run.text.strip():
                                pPr = para._p.get_or_add_pPr()
                                # Remove existing noWrap to avoid duplicates
                                for child in pPr.findall(qn('w:rPr')):
                                    for nw in child.findall(qn('w:noBreakHyphen')):
                                        child.remove(nw)
                                break

                        # Keep row with next row (prevent table splitting across pages)
                        if row_idx < num_rows - 1:
                            para.paragraph_format.keep_with_next = True

                    # Borders
                    # Determine styling based on position (Edge vs Inside)
                    is_first_row = (row_idx == 0)
                    is_last_row = (row_idx == num_rows - 1)
                    is_first_col = (cell_idx == 0)
                    is_last_col = (cell_idx == num_cols - 1)

                    border_kwargs = {}

                    # Top Border
                    if is_first_row:
                        border_kwargs['top'] = {'val': 'single', 'sz': border_outline_sz, 'color': border_outline_color_hex}
                    else:
                        if border_h != 'nil':
                            border_kwargs['top'] = {'val': border_h, 'sz': '4', 'color': border_h_color_hex}

                    # Bottom Border
                    if is_last_row:
                        border_kwargs['bottom'] = {'val': 'single', 'sz': border_outline_sz, 'color': border_outline_color_hex}
                    else:
                        if border_h != 'nil':
                            border_kwargs['bottom'] = {'val': border_h, 'sz': '4', 'color': border_h_color_hex}

                    # Left Border
                    if is_first_col:
                        border_kwargs['left'] = {'val': 'single', 'sz': border_outline_sz, 'color': border_outline_color_hex}
                    else:
                        if border_v != 'nil':
                            border_kwargs['left'] = {'val': border_v, 'sz': '4', 'color': border_v_color_hex}

                    # Right Border
                    if is_last_col:
                        border_kwargs['right'] = {'val': 'single', 'sz': border_outline_sz, 'color': border_outline_color_hex}
                    else:
                        if border_v != 'nil':
                            border_kwargs['right'] = {'val': border_v, 'sz': '4', 'color': border_v_color_hex}

                    if border_kwargs:
                        set_cell_border(cell, **border_kwargs)

if __name__ == '__main__':
    app.run(debug=True, port=3000)
